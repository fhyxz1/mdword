from kivy.app import App
from kivy.core.text import LabelBase
from kivy.resources import resource_add_path, resource_find
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.popup import Popup
from kivy.uix.label import Label
from kivy.uix.button import Button
from kivy.uix.textinput import TextInput
from kivy.uix.scrollview import ScrollView
from kivy.lang import Builder
from kivy.properties import BooleanProperty, OptionProperty
from kivy.clock import Clock
import re
import os
import sys
import time
from kivy.utils import platform

# 修改Word处理库导入方式，在所有平台上尝试导入
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    
    # 定义常用颜色
    BLACK_COLOR = RGBColor(0, 0, 0)
    BLUE_COLOR = RGBColor(0, 0, 255)
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

# Android剪贴板处理
if platform == 'android':
    from jnius import autoclass
    PythonActivity = autoclass('org.kivy.android.PythonActivity')
    Context = autoclass('android.content.Context')
    ClipData = autoclass('android.content.ClipData')
    clipboard_service = PythonActivity.mActivity.getSystemService(Context.CLIPBOARD_SERVICE)

    def android_copy(text):
        clip = ClipData.newPlainText("text", text)
        clipboard_service.setPrimaryClip(clip)

    def android_paste():
        clip = clipboard_service.getPrimaryClip()
        if clip and clip.getItemCount() > 0:
            item = clip.getItemAt(0)
            return str(item.coerceToText(PythonActivity.mActivity))
        return ""
else:
    from kivy.core.clipboard import Clipboard
    def android_copy(text):
        Clipboard.copy(text)
    def android_paste():
        return Clipboard.paste()


# 字体处理
if hasattr(sys, '_MEIPASS'):
    resource_add_path(os.path.join(sys._MEIPASS, 'fonts'))
else:
    resource_add_path(os.path.abspath('./fonts'))

LabelBase.register('Roboto', resource_find('SourceHanSansSC-Regular-2.otf'))

Builder.load_string('''
#:kivy 2.0.0
#:import hex kivy.utils.get_color_from_hex

<MarkdownTool>:
    orientation: 'vertical'
    spacing: '8dp'
    padding: '8dp'

    BoxLayout:
        orientation: 'vertical'
        size_hint_y: 0.30
        Label:
            text: '输入区'
            size_hint_y: None
            height: '25dp'
            font_size: '14sp'
        ScrollView:
            TextInput:
                id: input_area
                text: ''
                hint_text: '在此输入或粘贴Markdown内容...'
                size_hint_y: None
                height: max(self.minimum_height, dp(200))
                background_color: hex('#FFFFFF')
                foreground_color: hex('#333333')
                font_size: '14sp'
                on_text: root.auto_process_and_update()

    BoxLayout:
        orientation: 'vertical'
        size_hint_y: 0.35
        spacing: '4dp'
        padding: '4dp'
        Label:
            text: '选项区'
            size_hint_y: None
            height: '25dp'
            font_size: '14sp'
        ScrollView:
            GridLayout:
                cols: 2
                size_hint_y: None
                height: self.minimum_height
                row_default_height: '32dp'
                row_force_default: True
                spacing: '8dp'
                padding: '4dp'
                CheckBox:
                    size_hint_x: None
                    width: '32dp'
                    active: root.remove_italic
                    on_active: root.remove_italic = self.active
                Label:
                    text: '去除斜体'
                    font_size: '12sp'
                CheckBox:
                    size_hint_x: None
                    width: '32dp'
                    active: root.remove_strikethrough
                    on_active: root.remove_strikethrough = self.active
                Label:
                    text: '删除线'
                    font_size: '12sp'
                CheckBox:
                    size_hint_x: None
                    width: '32dp'
                    active: root.remove_highlight
                    on_active: root.remove_highlight = self.active
                Label:
                    text: '去除高亮'
                    font_size: '12sp'
                CheckBox:
                    size_hint_x: None
                    width: '32dp'
                    active: root.remove_links
                    on_active: root.remove_links = self.active
                Label:
                    text: '去除链接'
                    font_size: '12sp'
                CheckBox:
                    size_hint_x: None
                    width: '32dp'
                    active: root.remove_unordered_list
                    on_active: root.remove_unordered_list = self.active
                Label:
                    text: '清洗无序列表'
                    font_size: '12sp'
                CheckBox:
                    size_hint_x: None
                    width: '32dp'
                    active: root.remove_ordered_list
                    on_active: root.remove_ordered_list = self.active
                Label:
                    text: '清洗有序列表'
                    font_size: '12sp'
                CheckBox:
                    size_hint_x: None
                    width: '32dp'
                    active: root.table_clean
                    on_active: root.table_clean = self.active
                Label:
                    text: '表格清洁'
                    font_size: '12sp'
                Label:
                    text: '表格转换:'
                    font_size: '12sp'
                Spinner:
                    id: table_spinner
                    text: root.table_conversion
                    values: ['无', '空格', '/t', ',']
                    font_size: '12sp'
                    size_hint_y: None
                    height: '32dp'
                    on_text: root.table_conversion = self.text
                CheckBox:
                    size_hint_x: None
                    width: '32dp'
                    active: root.table_to_word
                    on_active: root.table_to_word = self.active
                Label:
                    text: '表格转Word表格'
                    font_size: '12sp'

    BoxLayout:
        orientation: 'vertical'
        size_hint_y: 0.30
        Label:
            text: '输出区'
            size_hint_y: None
            height: '25dp'
            font_size: '14sp'
        BoxLayout:
            size_hint_y: None
            height: '32dp'
            spacing: '4dp'
            CustomButton:
                text: '清空输入'
                on_press: root.process_reset('input')
                font_size: '8sp'
                size_hint_x: 1
            CustomButton:
                text: '读取剪贴板'
                on_press: root.paste_from_clipboard()
                font_size: '9sp'
                size_hint_x: 1
            CustomButton:
                text: '清空输出'
                on_press: root.process_reset('output')
                font_size: '8sp'
                size_hint_x: 1
            CustomButton:
                text: '复制结果'
                on_press: root.copy_to_clipboard()
                font_size: '8sp'
                size_hint_x: 1
        
        BoxLayout:
            size_hint_y: None
            height: '32dp'
            spacing: '4dp'
            CustomButton:
                text: '导出为Word'
                on_press: root.export_to_word()
                font_size: '12sp'
                size_hint_x: 1
                background_color: hex('#3F51B5') if self.state == 'normal' else hex('#303F9F')
                # 按钮状态由代码控制
        ScrollView:
            TextInput:
                id: output_area
                text: ''
                hint_text: '处理后的纯净文本...'
                size_hint_y: None
                height: max(self.minimum_height, dp(200))
                background_color: hex('#F5F5F5')
                foreground_color: hex('#333333')
                font_size: '14sp'

<CustomButton@Button>:
    font_size: '12sp'
    background_normal: ''
    background_color: hex('#4CAF50') if self.state == 'normal' else hex('#45a049')
    color: hex('#FFFFFF')
    size_hint: (None, None)
    size: ('80dp', '32dp')
    padding: ('4dp', '4dp')
''')

# [剩余类定义保持不变...]

class MarkdownTool(BoxLayout):
    auto_process = BooleanProperty(True)
    remove_italic = BooleanProperty(False)
    remove_strikethrough = BooleanProperty(False)
    remove_highlight = BooleanProperty(False)
    remove_links = BooleanProperty(False)
    remove_unordered_list = BooleanProperty(False)
    remove_ordered_list = BooleanProperty(False)
    table_clean = BooleanProperty(False)
    table_conversion = OptionProperty("无", options=["无", "空格", "/t", ","])
    table_to_word = BooleanProperty(True)  # 是否将Markdown表格转换为Word表格

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.bind(
            remove_italic=lambda inst, val: self._option_changed(),
            remove_strikethrough=lambda inst, val: self._option_changed(),
            remove_highlight=lambda inst, val: self._option_changed(),
            remove_links=lambda inst, val: self._option_changed(),
            remove_unordered_list=lambda inst, val: self._option_changed(),
            remove_ordered_list=lambda inst, val: self._option_changed(),
            table_clean=lambda inst, val: self._option_changed(),
            table_conversion=lambda inst, val: self._option_changed(),
            table_to_word=lambda inst, val: self._option_changed()
        )
        
        # 注册窗口加载后的回调函数，用于设置导出按钮状态
        Clock.schedule_once(self._set_export_button_state)

    def _option_changed(self):
        if self.auto_process:
            self.process_markdown()

    def paste_from_clipboard(self):
        try:
            self.ids.input_area.text = android_paste().strip()
        except Exception as e:
            self.ids.output_area.text = f"剪贴板错误: {str(e)}"

    def auto_process_and_update(self):
        if self.auto_process:
            self.process_markdown()

    def process_markdown(self):
        try:
            text = self.ids.input_area.text
            text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r"''(.*?)''", r'\1', text)

            if self.remove_italic:
                text = re.sub(r'(?<!\*)\*(?!\*)(.*?)\*(?!\*)', r'\1', text)
                text = re.sub(r'(?<!_)_(?!_)(.*?)_(?!_)', r'\1', text)
            if self.remove_strikethrough:
                text = re.sub(r'~~(.*?)~~', r'\1', text)
            if self.remove_highlight:
                text = re.sub(r'==(.+?)==', r'\1', text)
            if self.remove_links:
                text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            if self.remove_unordered_list:
                text = re.sub(r'(?m)^\s*[-*+]\s+', '', text)
            if self.remove_ordered_list:
                text = re.sub(r'(?m)^\s*\d+\.\s+', '', text)

            if self.table_clean:
                text = re.sub(r'(?m)^\s*\|?[\s\-|]+\|?\s*$', '', text)
                text = text.replace("|", "")
            elif self.table_conversion != "无":
                lines = text.splitlines()
                processed_lines = []
                for line in lines:
                    if re.match(r'^\s*\|?[\s\-|]+\|?\s*$', line):
                        continue
                    line = line.strip()
                    if line.startswith("|"):
                        line = line[1:]
                    if line.endswith("|"):
                        line = line[:-1]
                    if self.table_conversion == "空格":
                        line = line.replace("|", "    ")
                    elif self.table_conversion == "/t":
                        line = line.replace("|", "\t")
                    elif self.table_conversion == ",":
                        line = line.replace("|", ",")
                    processed_lines.append(line)
                text = "\n".join(processed_lines)

            text = re.sub(r'(?m)^(?:\s*[-*_]{3,}\s*)$', '', text)
            self.ids.output_area.text = text.strip()
        except Exception as e:
            self.ids.output_area.text = f"处理错误: {str(e)}"

    def copy_to_clipboard(self):
        try:
            android_copy(self.ids.output_area.text)
        except Exception as e:
            self.ids.output_area.text = f"复制失败: {str(e)}"

    def process_reset(self, target):
        getattr(self.ids, f"{target}_area").text = ''

    def _extract_table(self, lines, start_index):
        """从Markdown文本中提取完整的表格并返回表格数据和新的索引"""
        table_data = []
        i = start_index
        has_header_separator = False
        
        while i < len(lines):
            line = lines[i].strip()
            
            # 如果行为空或不包含|，则表格结束
            if not line or '|' not in line:
                break
                
            # 检查是否为分隔行（包含 ----|----）
            if re.match(r'^\s*\|?[\s\-:|]+\|?\s*$', line):
                has_header_separator = True
                i += 1
                continue
                
            # 提取单元格数据
            cells = [cell.strip() for cell in line.strip('|').split('|')]
            table_data.append(cells)
            
            i += 1
        
        return table_data, has_header_separator, i
        
    def _process_table(self, doc, table_data, has_header):
        """将提取的表格数据转换为Word表格"""
        if not table_data or not table_data[0]:
            return
            
        # 创建表格
        rows_count = len(table_data)
        cols_count = max(len(row) for row in table_data)
        
        table = doc.add_table(rows=rows_count, cols=cols_count)
        table.style = 'Table Grid'  # 应用表格网格样式
        
        # 填充表格内容
        for i, row_data in enumerate(table_data):
            for j, cell_content in enumerate(row_data):
                if j < cols_count:  # 确保不超出列数
                    cell = table.cell(i, j)
                    # 处理单元格中的Markdown格式
                    para = cell.paragraphs[0]
                    self._add_text_with_bold(para, cell_content)
                    
                    # 如果有表头且这是第一行，设置加粗和居中
                    if has_header and i == 0:
                        for run in para.runs:
                            run.bold = True
                            run.font.name = '黑体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
        # 添加表格后的空行
        doc.add_paragraph()
    
    def _add_text_with_bold(self, paragraph, text):
        """处理文本，保留加粗格式"""
        # 查找所有加粗部分
        bold_parts = re.finditer(r'\*\*(.*?)\*\*', text)
        
        last_end = 0
        has_bold = False
        
        for match in bold_parts:
            has_bold = True
            start, end = match.span()
            # 添加加粗前的普通文本
            if start > last_end:
                normal_run = paragraph.add_run(text[last_end:start])
                normal_run.font.name = '宋体'
                # 添加东亚字体设置
                normal_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
            # 添加加粗文本 - 明确设置为加粗且使用黑体
            bold_run = paragraph.add_run(match.group(1))
            bold_run.bold = True  # 确保加粗
            bold_run.font.name = '黑体'  # 加粗文本使用黑体
            # 添加东亚字体设置
            bold_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            
            last_end = end
            
        # 如果没有加粗部分，直接添加整个文本
        if not has_bold and last_end == 0:
            # 处理其他可能的Markdown标记
            processed_text = self._process_markdown_inline(text)
            normal_run = paragraph.add_run(processed_text)
            normal_run.font.name = '宋体'
            # 添加东亚字体设置
            normal_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            return
            
        # 添加最后剩余的普通文本
        if last_end < len(text):
            # 移除可能存在的其他Markdown标记
            remaining_text = text[last_end:]
            processed_text = self._process_markdown_inline(remaining_text)
            
            normal_run = paragraph.add_run(processed_text)
            normal_run.font.name = '宋体'
            # 添加东亚字体设置
            normal_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
    def _process_markdown_inline(self, text):
        """处理内联Markdown标记"""
        processed_text = text
        
        # 移除斜体（如果启用了该选项）
        if self.remove_italic:
            processed_text = re.sub(r'(?<!\*)\*(?!\*)(.*?)\*(?!\*)', r'\1', processed_text)
            processed_text = re.sub(r'(?<!_)_(?!_)(.*?)_(?!_)', r'\1', processed_text)
            
        # 移除删除线（如果启用了该选项）
        if self.remove_strikethrough:
            processed_text = re.sub(r'~~(.*?)~~', r'\1', processed_text)
            
        # 移除高亮（如果启用了该选项）
        if self.remove_highlight:
            processed_text = re.sub(r'==(.+?)==', r'\1', processed_text)
            
        # 移除链接（如果启用了该选项）
        if self.remove_links:
            processed_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', processed_text)
            
        return processed_text
        
    def export_to_word(self):
        """导出为Word文档功能"""
        if not WORD_AVAILABLE:
            # 更有帮助的错误消息
            self.show_message_popup("无法导出Word文档: 缺少必要的库。\n请确保已安装 python-docx 和 lxml 库。")
            return
            
        if not self.ids.input_area.text.strip():
            self.show_message_popup("请先输入或粘贴Markdown内容")
            return

        try:
            # 针对不同平台使用不同的导出方法
            if platform == 'android':
                # 记录调试信息
                print("正在尝试Android SAF导出")
                self._export_with_android_saf()
            else:
                # 使用系统文件选择器
                try:
                    from plyer import filechooser
                    
                    # 设置默认文件名
                    default_file = os.path.join(os.path.expanduser("~"), "md导出文档.docx")
                    
                    # 打开保存文件对话框
                    filepath = filechooser.save_file(title="保存Markdown为Word文档", 
                                                   filters=[("Word文档", "*.docx")],
                                                   defaultextension=".docx",
                                                   path=os.path.dirname(default_file),
                                                   filename=os.path.basename(default_file))
                                                   
                    if filepath and len(filepath) > 0:
                        filepath = filepath[0]  # plyer返回的是列表
                        # 确保文件名有.docx扩展名
                        if not filepath.lower().endswith('.docx'):
                            filepath += '.docx'
                            
                        # 保存文档
                        self.save_word_document(filepath)
                except ImportError:
                    # 如果plyer不可用，则使用tkinter
                    if platform != 'android':
                        self._export_with_tkinter()
                    else:
                        self.show_message_popup("无法找到适合的文件选择器")
                
        except Exception as e:
            # 更详细的错误消息
            import traceback
            error_details = traceback.format_exc()
            print(f"导出错误详情: {error_details}")
            self.show_message_popup(f"Word导出失败: {str(e)}\n\n请确保应用有存储权限。")
    
    def _export_with_android_saf(self):
        """使用Android的Storage Access Framework API导出文档"""
        try:
            from jnius import autoclass, cast
            from android.runnable import run_on_ui_thread
            import threading
            
            # 检查并请求存储权限
            self._check_android_permissions()
            
            # 主要的Android类
            PythonActivity = autoclass('org.kivy.android.PythonActivity')
            Intent = autoclass('android.content.Intent')
            Environment = autoclass('android.os.Environment')
            Uri = autoclass('android.net.Uri')
            FileOutputStream = autoclass('java.io.FileOutputStream')
            ParcelFileDescriptor = autoclass('android.os.ParcelFileDescriptor')
            ContentResolver = autoclass('android.content.ContentResolver')
            
            # 定义文件类型和用于Word文档的MIME类型
            DOC_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
            # 创建一个临时文件来保存Word文档
            temp_dir = os.path.join(PythonActivity.mActivity.getCacheDir().getAbsolutePath(), "temp")
            if not os.path.exists(temp_dir):
                os.makedirs(temp_dir)
            temp_file = os.path.join(temp_dir, "temp_document.docx")
            
            # 首先保存到临时文件
            self.save_word_document(temp_file)
            
            # 定义结果处理的回调
            result_handler = []
            
            # 定义请求码
            REQUEST_CODE_CREATE_FILE = 43
            
            # 定义SAF活动结果回调
            class ActivityResultListener:
                def __init__(self, callback):
                    self.callback = callback
                
                def onActivityResult(self, requestCode, resultCode, data):
                    if requestCode == REQUEST_CODE_CREATE_FILE:
                        # 活动结果代码定义
                        RESULT_OK = -1
                        if resultCode == RESULT_OK and data:
                            uri = data.getData()
                            self.callback(uri)
                        else:
                            # 用户取消或出错
                            pass
                    return True
            
            # 处理URI的回调
            def handle_uri(uri):
                try:
                    # 定义请求码
                    REQUEST_CODE_GET_CONTENT = 44
                    
                    # 定义SAF活动结果回调
                    class UriContentListener:
                        def __init__(self, callback):
                            self.callback = callback
                        
                        def onActivityResult(self, requestCode, resultCode, data):
                            if requestCode == REQUEST_CODE_GET_CONTENT:
                                # 活动结果代码定义
                                RESULT_OK = -1
                                if resultCode == RESULT_OK and data:
                                    content_uri = data.getData()
                                    self.callback(content_uri)
                                else:
                                    # 用户取消或出错
                                    pass
                            return True
                    
                    # 处理内容URI的回调
                    def handle_content_uri(content_uri):
                        try:
                            # 获取文件输出流
                            content_resolver = PythonActivity.mActivity.getContentResolver()
                            descriptor = content_resolver.openFileDescriptor(uri, "w")
                            if descriptor:
                                file_descriptor = descriptor.detachFd()
                                output_stream = ParcelFileDescriptor.AutoCloseOutputStream(
                                    ParcelFileDescriptor.adoptFd(file_descriptor))
                                
                                # 读取临时文件并写入到新创建的文件
                                with open(temp_file, 'rb') as f:
                                    data = f.read()
                                    output_stream.write(data)
                                    output_stream.flush()
                                    output_stream.close()
                                
                                # 显示成功消息
                                def show_success():
                                    options_summary = []
                                    if self.remove_italic:
                                        options_summary.append("去除斜体")
                                    if self.remove_strikethrough:
                                        options_summary.append("去除删除线")
                                    if self.remove_highlight:
                                        options_summary.append("去除高亮")
                                    if self.remove_links:
                                        options_summary.append("去除链接")
                                    if self.remove_unordered_list:
                                        options_summary.append("清洗无序列表")
                                    if self.remove_ordered_list:
                                        options_summary.append("清洗有序列表")
                                    if self.table_clean:
                                        options_summary.append("表格清洁")
                                    elif self.table_conversion != "无":
                                        options_summary.append(f"表格转换: {self.table_conversion}")
                                    if self.table_to_word and not self.table_clean and self.table_conversion == "无":
                                        options_summary.append("表格转为Word表格")
                                        
                                    options_text = "，".join(options_summary) if options_summary else "保留所有格式"
                                    success_message = f"文档已成功导出\n\n应用的处理选项: {options_text}"
                                    
                                    self.show_message_popup(success_message)
                                    
                                # 在主线程上显示成功消息
                                Clock.schedule_once(lambda dt: show_success(), 0)
                                
                        except Exception as e:
                            # 在主线程上显示错误消息
                            def show_error():
                                self.show_message_popup(f"保存文件时出错: {str(e)}")
                            Clock.schedule_once(lambda dt: show_error(), 0)
                    
                    # 创建结果监听器并注册
                    content_listener = UriContentListener(handle_content_uri)
                    
                    # 创建一个列表来保持监听器的引用
                    global uri_content_handlers
                    if not hasattr(sys.modules[__name__], 'uri_content_handlers'):
                        uri_content_handlers = []
                    uri_content_handlers.append(content_listener)
                    
                    # 注册活动结果回调
                    PythonActivity.mActivity.registerActivityResultListener(content_listener)
                    
                    # 在UI线程上启动获取内容Intent
                    @run_on_ui_thread
                    def get_content():
                        # 创建Intent以获取内容
                        intent = Intent(Intent.ACTION_GET_CONTENT)
                        intent.addCategory(Intent.CATEGORY_OPENABLE)
                        intent.setType("*/*")
                        
                        # 启动活动
                        PythonActivity.mActivity.startActivityForResult(intent, REQUEST_CODE_GET_CONTENT)
                    
                    # 执行获取内容Intent
                    get_content()
                    
                except Exception as e:
                    # 在主线程上显示错误消息
                    def show_error():
                        self.show_message_popup(f"处理URI时出错: {str(e)}")
                    Clock.schedule_once(lambda dt: show_error(), 0)
            
            # 创建结果监听器并注册
            listener = ActivityResultListener(handle_uri)
            result_handler.append(listener)
            
            # 注册活动结果回调
            PythonActivity.mActivity.registerActivityResultListener(listener)
            
            # 在UI线程上启动文件创建Intent
            @run_on_ui_thread
            def create_file():
                current_time = autoclass('java.lang.System').currentTimeMillis()
                default_filename = f"md导出文档_{current_time}.docx"
                
                # 创建Intent以创建文件
                intent = Intent(Intent.ACTION_CREATE_DOCUMENT)
                intent.addCategory(Intent.CATEGORY_OPENABLE)
                intent.setType(DOC_MIME_TYPE)
                intent.putExtra(Intent.EXTRA_TITLE, default_filename)
                
                # 启动活动
                PythonActivity.mActivity.startActivityForResult(intent, REQUEST_CODE_CREATE_FILE)
            
            # 执行文件创建Intent
            create_file()
            
        except Exception as e:
            self.show_message_popup(f"Android导出失败: {str(e)}")
            
    def _check_android_permissions(self):
        """检查并请求Android存储权限"""
        if platform == 'android':
            try:
                from android.permissions import request_permissions, Permission
                from android.permissions import check_permission
                
                # 检查是否有存储权限
                if not check_permission(Permission.WRITE_EXTERNAL_STORAGE):
                    # 请求存储权限
                    request_permissions([
                        Permission.READ_EXTERNAL_STORAGE,
                        Permission.WRITE_EXTERNAL_STORAGE
                    ])
                    # 提示用户已请求权限
                    self.show_message_popup("已请求存储权限，请授予权限后重试")
                    return False
                return True
            except Exception as e:
                print(f"权限检查错误: {e}")
                return False
        return True

    def _export_with_tkinter(self):
        """使用tkinter的文件选择器导出Word文档"""
        try:
            from tkinter import Tk, filedialog
            import tkinter
            
            # 创建一个隐藏的Tkinter根窗口
            try:
                root = Tk()
                root.withdraw()  # 隐藏Tkinter窗口
                
                # 使用Tkinter的文件对话框获取保存路径
                filepath = filedialog.asksaveasfilename(
                    title="保存Markdown为Word文档",
                    defaultextension=".docx",
                    filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")],
                    initialdir=os.path.expanduser("~/Documents"),
                    initialfile="md导出文档.docx"
                )
                
                # 如果用户取消，filepath会是空字符串
                if filepath:
                    # 确保文件名有.docx扩展名
                    if not filepath.lower().endswith('.docx'):
                        filepath += '.docx'
                        
                    # 保存文档
                    self.save_word_document(filepath)
                    
            except Exception as e:
                self.show_message_popup(f"Tkinter导出对话框错误: {e}")
            finally:
                try:
                    # 销毁Tkinter窗口
                    root.destroy()
                except:
                    pass
        except ImportError:
            self.show_message_popup("无法导入tkinter模块")
            
    def save_word_document(self, filepath):
        """将Markdown内容保存为Word文档"""
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置基本样式
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal'].font.size = Pt(12)
            # 确保中文字体名称可以识别
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 设置标题样式：黑体、黑色、加粗
            for i in range(1, 10):  # Word支持9级标题
                heading_name = f'Heading {i}'
                if heading_name in doc.styles:
                    heading_style = doc.styles[heading_name]
                    # 明确设置黑体字
                    heading_style.font.name = '黑体' 
                    # 确保中文字体名称可以识别，有些环境可能需要英文字体名
                    heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    heading_style.font.bold = True  # 加粗
                    heading_style.font.italic = False  # 不斜体
                    # 确保标题字体颜色为黑色
                    heading_style.font.color.rgb = BLACK_COLOR  # 使用常量
            
            # 获取原始Markdown文本
            md_text = self.ids.input_area.text
            
            # 处理Markdown文本，保留基本结构
            lines = md_text.splitlines()
            
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                
                # 跳过空行
                if not line:
                    doc.add_paragraph()
                    i += 1
                    continue
                
                # 处理标题 - 正确对应Markdown标题级别和Word标题级别
                heading_match = re.match(r'^(#+)\s+(.*)', line)
                if heading_match:
                    # 获取#的数量作为标题级别
                    hash_count = len(heading_match.group(1))
                    # Markdown中#对应一级标题，##对应二级标题，以此类推
                    level = hash_count
                    text = heading_match.group(2)
                    
                    # 添加对应级别的标题
                    heading = doc.add_heading(text, level=level)
                    
                    # 确保标题使用黑体，加粗，不斜体，颜色为黑色
                    for run in heading.runs:
                        run.font.name = '黑体'
                        # 确保中文字体名称可以识别
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        run.bold = True  # 确保加粗
                        run.italic = False
                        run.font.color.rgb = BLACK_COLOR  # 使用常量设置为黑色
                    
                    i += 1
                    continue
                
                # 处理无序列表
                list_match = re.match(r'^\s*([-*+])\s+(.*)', line)
                if list_match and not self.remove_unordered_list:
                    list_item = doc.add_paragraph(style='List Bullet')
                    
                    # 提取列表项文本并处理其中的加粗
                    text = list_match.group(2)
                    self._add_text_with_bold(list_item, text)
                    
                    i += 1
                    continue
                elif list_match and self.remove_unordered_list:
                    # 如果设置了移除无序列表，则作为普通段落处理
                    para = doc.add_paragraph()
                    self._add_text_with_bold(para, list_match.group(2))
                    i += 1
                    continue
                
                # 处理有序列表
                ordered_list_match = re.match(r'^\s*(\d+\.)\s+(.*)', line)
                if ordered_list_match and not self.remove_ordered_list:
                    list_item = doc.add_paragraph(style='List Number')
                    self._add_text_with_bold(list_item, ordered_list_match.group(2))
                    i += 1
                    continue
                elif ordered_list_match and self.remove_ordered_list:
                    # 如果设置了移除有序列表，则作为普通段落处理
                    para = doc.add_paragraph()
                    self._add_text_with_bold(para, ordered_list_match.group(2))
                    i += 1
                    continue
                
                # 处理表格（简化处理，表格转换为制表符）
                if '|' in line and not line.startswith('>'):
                    # 如果启用了表格转Word表格选项并且不是清洁模式
                    if self.table_to_word and not self.table_clean and self.table_conversion == "无":
                        # 提取完整表格
                        table_data, has_header, i = self._extract_table(lines, i)
                        # 处理表格为Word表格
                        self._process_table(doc, table_data, has_header)
                        continue
                    
                    # 检查是否是表格分隔行
                    if re.match(r'^\s*\|?[\s\-:|]+\|?\s*$', line):
                        i += 1
                        continue
                    
                    # 按照表格转换设置处理
                    cells = [cell.strip() for cell in line.strip('|').split('|')]
                    
                    # 根据表格转换选项处理
                    if self.table_clean:
                        # 如果启用了table_clean，直接处理为简单文本，使用空格分隔
                        text = " ".join(cells)
                    elif self.table_conversion == "无":
                        # 默认使用空格
                        text = " ".join(cells)
                    elif self.table_conversion == "空格":
                        text = "    ".join(cells)
                    elif self.table_conversion == "/t":
                        text = "\t".join(cells)
                    elif self.table_conversion == ",":
                        text = ",".join(cells)
                    else:
                        # 默认使用空格
                        text = " ".join(cells)
                    
                    para = doc.add_paragraph()
                    self._add_text_with_bold(para, text)
                    i += 1
                    continue
                
                # 处理普通段落和其他格式
                para = doc.add_paragraph()
                
                # 处理文本，保留加粗格式
                self._add_text_with_bold(para, line)
                i += 1
            
            # 保存Word文档
            doc.save(filepath)
            
            # 生成处理选项的摘要
            options_summary = []
            if self.remove_italic:
                options_summary.append("去除斜体")
            if self.remove_strikethrough:
                options_summary.append("去除删除线")
            if self.remove_highlight:
                options_summary.append("去除高亮")
            if self.remove_links:
                options_summary.append("去除链接")
            if self.remove_unordered_list:
                options_summary.append("清洗无序列表")
            if self.remove_ordered_list:
                options_summary.append("清洗有序列表")
            if self.table_clean:
                options_summary.append("表格清洁")
            elif self.table_conversion != "无":
                options_summary.append(f"表格转换: {self.table_conversion}")
            if self.table_to_word and not self.table_clean and self.table_conversion == "无":
                options_summary.append("表格转为Word表格")
                
            options_text = "，".join(options_summary) if options_summary else "保留所有格式"
            success_message = f"文档已成功保存至:\n{filepath}\n\n应用的处理选项: {options_text}"
            
            self.show_message_popup(success_message)
            
        except Exception as e:
            error_msg = f"Word导出失败: {str(e)}"
            self.show_message_popup(error_msg)
            
    def show_message_popup(self, message):
        """显示消息弹窗"""
        content = BoxLayout(orientation='vertical', spacing=10, padding=10)
        popup = Popup(title='提示', content=content, size_hint=(0.8, 0.5))
        
        # 添加消息标签
        content.add_widget(Label(text=message))
        
        # 添加确认按钮
        btn = Button(text='确定', size_hint=(1, 0.3))
        btn.bind(on_press=popup.dismiss)
        content.add_widget(btn)
        
        popup.open()

    def _set_export_button_state(self, dt):
        """设置导出Word按钮的状态"""
        try:
            # 查找导出按钮
            export_button = None
            for child in self.children:
                if isinstance(child, BoxLayout) and child.size_hint_y == 0.30:  # 输出区
                    for subchild in child.children:
                        if isinstance(subchild, BoxLayout) and subchild.height == 32:  # 按钮行
                            if len(subchild.children) > 0:
                                button = subchild.children[0]
                                if hasattr(button, 'text') and button.text == '导出为Word':
                                    export_button = button
                                    break
            
            # 设置按钮状态
            if export_button:
                export_button.disabled = not WORD_AVAILABLE
                if not WORD_AVAILABLE:
                    export_button.opacity = 0.5  # 半透明表示不可用
                else:
                    export_button.opacity = 1  # 完全不透明表示可用
        except Exception as e:
            print(f"设置导出按钮状态时出错: {e}")

class MarkdownApp(App):
    def build(self):
        self.title = 'mdword'
        # 应用移动端优化
        if platform == 'android':
            # Android端优化
            from kivy.core.window import Window
            from kivy.metrics import dp
            
            # 设置深色状态栏
            try:
                from android.runnable import run_on_ui_thread
                from jnius import autoclass
                
                WindowManager = autoclass('android.view.WindowManager$LayoutParams')
                activity = autoclass('org.kivy.android.PythonActivity').mActivity
                
                @run_on_ui_thread
                def set_statusbar():
                    # 设置状态栏颜色
                    window = activity.getWindow()
                    # 设置为深色状态栏
                    window.setStatusBarColor(0xFF4CAF50)
                    # 设置状态栏文字为浅色
                    window.getDecorView().setSystemUiVisibility(0)
                
                set_statusbar()
            except Exception as e:
                print(f"设置状态栏时出错: {e}")
            
            # 适配全面屏
            Window.softinput_mode = 'below_target'
            
        # 返回主界面
        return MarkdownTool()

if __name__ == '__main__':
    MarkdownApp().run()