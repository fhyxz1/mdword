from docx.oxml import parse_xml
from kivy.app import App
from kivy.core.text import LabelBase
from kivy.resources import resource_add_path, resource_find
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.popup import Popup
from kivy.uix.label import Label
from kivy.uix.button import Button
from kivy.uix.textinput import TextInput
from kivy.core.window import Window
from kivy.lang import Builder
from kivy.properties import BooleanProperty, OptionProperty
from kivy.clock import Clock
from kivy.uix.filechooser import FileChooserListView
import re
import threading
import pyperclip  # 更可靠的剪贴板库
import keyboard  # 全局快捷键支持
from pystray import Icon, Menu, MenuItem  # 系统托盘支持
from PIL import Image, ImageDraw  # 图标处理
import os
import sys
import time
from docx import Document  # Word文档处理库
from docx.shared import Pt, Inches, RGBColor  # 重新导入RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX  # 确保正确导入高亮颜色常量
from docx.enum.style import WD_STYLE_TYPE  # 添加样式类型导入
from docx.oxml.ns import qn, nsdecls  # 修正qn的导入路径

BLACK_COLOR = RGBColor(0, 0, 0)  # 黑色
BLUE_COLOR = RGBColor(0, 0, 255)  # 蓝色
GRAY_COLOR = RGBColor(128, 128, 128)  # 灰色
YELLOW_HIGHLIGHT = 7  # 黄色高亮的索引值

# Windows API 文件选择对话框支持
try:
    import win32gui
    import win32con
    import ctypes
    
    # 判断是否为Windows系统
    is_windows = sys.platform.startswith('win')
except ImportError:
    is_windows = False

"""
@Version: v1.2
@Author: ylab
@Date: 2025/3/25
@Update: 新增输出为Word文档功能
"""

if hasattr(sys, '_MEIPASS'):
    # 打包后的字体目录在 _MEIPASS/fonts 下
    resource_add_path(os.path.join(sys._MEIPASS, 'fonts'))
else:
    resource_add_path(os.path.abspath('./fonts'))

LabelBase.register('Roboto', resource_find('SourceHanSansSC-Regular-2.otf'))

Builder.load_string('''
#:kivy 2.0.0
#:import hex kivy.utils.get_color_from_hex

<MarkdownTool>:
    orientation: 'horizontal'
    spacing: '10sp'
    padding: '10sp'

    BoxLayout:
        orientation: 'vertical'
        size_hint_x: 0.4
        spacing: '5sp'

        BoxLayout:
            size_hint_y: 0.1
            spacing: '5sp'
            CustomButton:
                text: '清空输入'
                on_press: root.process_reset('input')
            CustomButton:
                text: '读取剪贴板'
                on_press: root.paste_from_clipboard()

        TextInput:
            id: input_area
            hint_text: '在此输入或粘贴Markdown内容...'
            background_color: hex('#FFFFFF')
            foreground_color: hex('#333333')
            on_text: root.auto_process_and_update()

    BoxLayout:
        orientation: 'vertical'
        size_hint_x: 0.2
        spacing: '15sp'
        padding: '10sp'

        # 处理选项区域
        BoxLayout:
            orientation: 'vertical'
            spacing: '5sp'

            BoxLayout:
                size_hint_y: None
                height: '30sp'
                CheckBox:
                    active: root.remove_italic
                    on_active: root.remove_italic = self.active
                Label:
                    text: '去除斜体'

            BoxLayout:
                size_hint_y: None
                height: '30sp'
                CheckBox:
                    active: root.remove_strikethrough
                    on_active: root.remove_strikethrough = self.active
                Label:
                    text: '删除线'

            BoxLayout:
                size_hint_y: None
                height: '30sp'
                CheckBox:
                    active: root.remove_highlight
                    on_active: root.remove_highlight = self.active
                Label:
                    text: '去除高亮'

            BoxLayout:
                size_hint_y: None
                height: '30sp'
                CheckBox:
                    active: root.remove_links
                    on_active: root.remove_links = self.active
                Label:
                    text: '去除链接'

            # 分开清洗无序列表与有序列表
            BoxLayout:
                size_hint_y: None
                height: '30sp'
                CheckBox:
                    active: root.remove_unordered_list
                    on_active: root.remove_unordered_list = self.active
                Label:
                    text: '清洗无序列表'
            BoxLayout:
                size_hint_y: None
                height: '30sp'
                CheckBox:
                    active: root.remove_ordered_list
                    on_active: root.remove_ordered_list = self.active
                Label:
                    text: '清洗有序列表'

            # 表格清洁复选框
            BoxLayout:
                size_hint_y: None
                height: '30sp'
                CheckBox:
                    active: root.table_clean
                    on_active: root.table_clean = self.active
                Label:
                    text: '表格清洁'

            # 表格转换下拉列表
            BoxLayout:
                size_hint_y: None
                height: '30sp'
                Label:
                    text: '表格转换:'
                Spinner:
                    id: table_spinner
                    text: root.table_conversion
                    values: ['无', '空格', '/t', ',']
                    on_text: root.table_conversion = self.text
            
            # 表格转为Word表格选项
            BoxLayout:
                size_hint_y: None
                height: '30sp'
                CheckBox:
                    active: root.table_to_word
                    on_active: root.table_to_word = self.active
                Label:
                    text: '表格转Word表格'

            # Word文档导出部分
            Label:
                text: 'Word导出选项:'
                size_hint_y: None
                height: '30sp'
                
            # Word导出按钮
            BoxLayout:
                size_hint_y: None
                height: '40sp'
                CustomButton:
                    text: '导出为Word'
                    on_press: root.export_to_word()

    BoxLayout:
        orientation: 'vertical'
        size_hint_x: 0.4
        spacing: '5sp'

        BoxLayout:
            size_hint_y: 0.1
            spacing: '5sp'
            CustomButton:
                text: '清空输出'
                on_press: root.process_reset('output')
            CustomButton:
                text: '复制结果'
                on_press: root.copy_to_clipboard()

        TextInput:
            id: output_area
            hint_text: '处理后的纯净文本...'
            background_color: hex('#F5F5F5')
            foreground_color: hex('#333333')
            font_name: 'Roboto'

<FileChooserPopup>:
    size_hint: 0.8, 0.8
    title: '选择保存Word文档的位置'
    BoxLayout:
        orientation: 'vertical'
        BoxLayout:
            orientation: 'vertical'
            size_hint_y: 0.15
            padding: '10sp'
            spacing: '5sp'
            
            BoxLayout:
                size_hint_y: None
                height: '30sp'
                Label:
                    text: '文件名:'
                    size_hint_x: 0.3
                TextInput:
                    id: filename_input
                    text: 'md导出文档.docx'
                    multiline: False
                    
            Label:
                text: '请选择保存位置:'
                size_hint_y: None
                height: '20sp'
                halign: 'left'
                text_size: self.size
        
        FileChooserListView:
            id: filechooser
            filters: ['*.docx']
            
        BoxLayout:
            size_hint_y: None
            height: '40sp'
            spacing: '5sp'
            padding: '10sp'
            CustomButton:
                text: '取消'
                on_press: root.dismiss()
            CustomButton:
                text: '保存'
                on_press: root.save(filechooser.path, filechooser.selection)

<CustomButton@Button>:
    font_size: '14sp'
    background_normal: ''
    background_color: hex('#4CAF50') if self.state == 'normal' else hex('#45a049')
    color: hex('#FFFFFF')
    size_hint_y: None
    height: '40sp'
''')

class FileChooserPopup(Popup):
    def __init__(self, save_callback, **kwargs):
        # 首先设置默认初始属性
        self.save_callback = save_callback
        # 调用父类的__init__方法
        super().__init__(**kwargs)
        # 设置文件选择器的初始路径
        self.ids.filechooser.path = os.path.expanduser('~\\Documents')
        
    def on_open(self):
        """弹窗打开时的初始化"""
        super().on_open()
        # 当弹窗打开时，设置焦点到文件名输入框
        Clock.schedule_once(lambda dt: self.ids.filename_input.focus, 0.1)

    def save(self, path, selection):
        """保存Word文档"""
        # 获取用户输入的文件名
        filename = self.ids.filename_input.text.strip()
        
        # 如果没有输入文件名，使用默认名称
        if not filename:
            filename = 'md导出文档.docx'
        
        # 确保文件名有.docx扩展名
        if not filename.lower().endswith('.docx'):
            filename += '.docx'
            
        # 构建完整的文件路径
        filepath = os.path.join(path, filename)
            
        # 调用回调函数保存文件
        self.save_callback(filepath)
        self.dismiss()

class MarkdownTool(BoxLayout):
    auto_process = BooleanProperty(True)
    remove_italic = BooleanProperty(False)
    remove_strikethrough = BooleanProperty(False)
    remove_highlight = BooleanProperty(False)
    remove_links = BooleanProperty(False)
    remove_unordered_list = BooleanProperty(False)
    remove_ordered_list = BooleanProperty(False)

    # 表格处理相关：清洁和转换选项
    table_clean = BooleanProperty(False)
    table_conversion = OptionProperty("无", options=["无", "空格", "/t",","])
    table_to_word = BooleanProperty(True)  # 是否将Markdown表格转换为Word表格

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self._keyboard = Window.request_keyboard(None, self)
        # 绑定选项变化时动态更新
        self.bind(remove_italic=lambda inst, val: self._option_changed())
        self.bind(remove_strikethrough=lambda inst, val: self._option_changed())
        self.bind(remove_highlight=lambda inst, val: self._option_changed())
        self.bind(remove_links=lambda inst, val: self._option_changed())
        self.bind(remove_unordered_list=lambda inst, val: self._option_changed())
        self.bind(remove_ordered_list=lambda inst, val: self._option_changed())
        self.bind(table_clean=lambda inst, val: self._option_changed())
        self.bind(table_conversion=lambda inst, val: self._option_changed())
        self.bind(table_to_word=lambda inst, val: self._option_changed())

    def _option_changed(self):
        if self.auto_process:
            self.process_markdown()

    def paste_from_clipboard(self):
        try:
            self.ids.input_area.text = pyperclip.paste().strip()
        except Exception as e:
            self.ids.output_area.text = f"剪贴板错误: {str(e)}"

    def auto_process_and_update(self):
        if self.auto_process:
            self.process_markdown()

    def process_markdown(self):
        start_time = time.perf_counter()
        try:
            text = self.ids.input_area.text

            # 移除 Markdown 标题
            text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)

            # 移除加粗语法 **text**
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)

            # 默认清洗功能：移除 ''md样式（两个单引号包裹）的文本标记
            text = re.sub(r"''(.*?)''", r'\1', text)

            # 添加对md``情况的默认处理（保留内容，去除格式）
            text = re.sub(r"md``(.*?)``", r'\1', text)

            # 改进：处理嵌套格式和多行情况的md``格式
            # 使用非贪婪匹配并支持跨行匹配
            text = re.sub(r"md``(.*?)``", r'\1', text, flags=re.DOTALL)

            # 根据选项去除斜体（处理 *text* 和 _text_ ）
            if self.remove_italic:
                text = re.sub(r'(?<!\*)\*(?!\*)(.*?)\*(?!\*)', r'\1', text)
                text = re.sub(r'(?<!_)_(?!_)(.*?)_(?!_)', r'\1', text)

            # 根据选项去除删除线语法 ~~text~~
            if self.remove_strikethrough:
                text = re.sub(r'~~(.*?)~~', r'\1', text)

            # 根据选项去除高亮语法（例如 ==text==）
            if self.remove_highlight:
                text = re.sub(r'==(.+?)==', r'\1', text)

            # 根据选项去除链接（Markdown 格式链接）
            if self.remove_links:
                text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

            # 列表样式清洁：分开处理无序列表和有序列表
            if self.remove_unordered_list:
                text = re.sub(r'(?m)^\s*[-*+]\s+', '', text)
            if self.remove_ordered_list:
                text = re.sub(r'(?m)^\s*\d+\.\s+', '', text)

            # 表格处理
            if self.table_clean:
                # 增强的清洁模式：处理对齐表格，移除分隔行、对齐标记和所有管道符
                lines = text.splitlines()
                processed_lines = []
                in_table = False

                for line in lines:
                    # 检测是否为表格分隔行（可能包含对齐标记 :-等）
                    if re.match(r'^\s*\|?[\s\-:|]+\|?\s*$', line):
                        in_table = True
                        continue  # 跳过分隔行

                    # 如果在表格中且行包含管道符
                    if in_table and '|' in line:
                        # 去除首尾管道符及空白
                        line = line.strip()
                        if line.startswith("|"):
                            line = line[1:]
                        if line.endswith("|"):
                            line = line[:-1]

                        # 移除列对齐标记（如 :--- 或 ---: 或 :--:）
                        line = re.sub(r':?-{3,}:?', '', line)

                        # 移除所有剩余的管道符
                        line = line.replace("|", "")

                        # 移除可能的多余空格
                        line = ' '.join(line.split())

                        processed_lines.append(line)
                    else:
                        if in_table:
                            in_table = False  # 表格结束
                        processed_lines.append(line)

                text = "\n".join(processed_lines)

            elif self.table_conversion != "无":
                # 分行处理：先将文本按行分割，再逐行清理
                lines = text.splitlines()
                processed_lines = []
                for line in lines:
                    # 跳过分隔行（包括对齐标记行）
                    if re.match(r'^\s*\|?[\s\-:|]+\|?\s*$', line):
                        continue
                    # 去除首尾可能存在的管道符及空白
                    line = line.strip()
                    if line.startswith("|"):
                        line = line[1:]
                    if line.endswith("|"):
                        line = line[:-1]
                    # 按选项转换中间的管道符
                    if self.table_conversion == "空格":
                        line = line.replace("|", "    ")
                    elif self.table_conversion == "/t":
                        line = line.replace("|", "\t")
                    elif self.table_conversion == ",":
                        line = line.replace("|", ",")
                    processed_lines.append(line)
                text = "\n".join(processed_lines)

            # 默认去除 Markdown 分割线（如 ---、***、___ 独占一行）
            text = re.sub(r'(?m)^(?:\s*[-*_]{3,}\s*)$', '', text)

            self.ids.output_area.text = text.strip()
        except Exception as e:
            self.ids.output_area.text = f"处理错误: {str(e)}"

    def copy_to_clipboard(self):
        try:
            pyperclip.copy(self.ids.output_area.text)
        except Exception as e:
            self.ids.output_area.text = f"复制失败: {str(e)}"

    def process_reset(self, target):
        getattr(self.ids, f"{target}_area").text = ''

    def export_to_word(self):
        """导出为Word文档功能"""
        if not self.ids.input_area.text.strip():
            self.show_message_popup("请先输入或粘贴Markdown内容")
            return

        try:
            # 尝试更安全的导出方式
            self.export_to_word_simple()
        except Exception as e:
            # 如果简化导出失败，回退到常规文件选择器
            print(f"简化导出失败: {e}, 类型: {type(e)}")
            # 回退到Kivy文件选择器
            popup = FileChooserPopup(self.save_word_document)
            popup.open()
            
    def export_to_word_simple(self):
        """简化版的Word导出功能，避免使用Windows API"""
        from tkinter import Tk, filedialog
        import tkinter
        
        # 创建一个隐藏的Tkinter根窗口
        try:
            root = Tk()
            root.withdraw()  # 隐藏Tkinter窗口
            
            # 使用Tkinter的文件对话框获取保存路径
            filepath = filedialog.asksaveasfilename(
                title="保存Markdown为Word文档",
                defaultextension=".docx",
                filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")],
                initialdir=os.path.expanduser("~/Documents"),
                initialfile="md导出文档.docx"
            )
            
            # 如果用户取消，filepath会是空字符串
            if filepath:
                # 确保文件名有.docx扩展名
                if not filepath.lower().endswith('.docx'):
                    filepath += '.docx'
                    
                # 保存文档
                self.save_word_document_simple(filepath)
                
        except Exception as e:
            print(f"Tkinter导出对话框错误: {e}")
            raise e  # 重新抛出异常，让调用者处理
        finally:
            try:
                # 销毁Tkinter窗口
                root.destroy()
            except:
                pass

    def save_word_document_simple(self, filepath):
        """简化版的Word导出功能，保留基本格式"""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            # 创建Word文档
            doc = Document()

            # 设置基本样式
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal'].font.size = Pt(12)
            # 确保中文字体名称可以识别
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 创建代码块样式（如果不存在）
            if 'Code' not in doc.styles:
                code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
                code_style.font.name = 'Courier New'
                code_style.font.size = Pt(10)
                code_style.paragraph_format.left_indent = Inches(0.5)
                code_style.paragraph_format.right_indent = Inches(0.5)
                code_style.paragraph_format.space_before = Pt(6)
                code_style.paragraph_format.space_after = Pt(6)
                
                # 设置底纹
                code_style_element = code_style._element
                if not code_style_element.rPr:
                    code_style_element.get_or_add_rPr()
                    
                # 为整个样式添加背景色
                shd_str = '<w:shd {}/>'.format(nsdecls('w'))
                shd = parse_xml(shd_str)
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F5F5F5')  # 浅灰色背景
                code_style_element.rPr.append(shd)

            # 设置标题样式：黑体、黑色、加粗
            for i in range(1, 10):  # Word支持9级标题
                heading_name = f'Heading {i}'
                if heading_name in doc.styles:
                    heading_style = doc.styles[heading_name]
                    # 明确设置黑体字
                    heading_style.font.name = '黑体'
                    # 确保中文字体名称可以识别，有些环境可能需要英文字体名
                    heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    heading_style.font.bold = True  # 加粗
                    heading_style.font.italic = False  # 不斜体
                    # 确保标题字体颜色为黑色
                    heading_style.font.color.rgb = BLACK_COLOR  # 使用常量

            # 获取原始Markdown文本
            md_text = self.ids.input_area.text

            # 处理Markdown文本，保留基本结构
            lines = md_text.splitlines()

            i = 0
            in_code_block = False
            code_block_content = []

            while i < len(lines):
                line = lines[i].strip()

                # 跳过空行
                if not line:
                    doc.add_paragraph()
                    i += 1
                    continue

                # 处理代码块开始和结束标记 ```
                if line.startswith("```"):
                    # 如果是代码块的开始
                    if not in_code_block:
                        in_code_block = True
                        code_block_content = []
                        # 提取语言信息
                        language = ""
                        if len(line) > 3:
                            language = line[3:].strip()
                        i += 1  # 跳过开始标记行
                        continue
                    else:
                        # 这是代码块的结束
                        in_code_block = False
                        
                        # 创建一个表格单元格来容纳代码，确保底纹能完整显示
                        code_table = doc.add_table(rows=1, cols=1)
                        code_table.autofit = True
                        code_table.style = 'Table Grid'
                        
                        # 移除表格边框
                        for cell in code_table._cells:
                            tcPr = cell._tc.get_or_add_tcPr()
                            tcBorders = parse_xml(r'<w:tcBorders %s><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>' % nsdecls('w'))
                            tcPr.append(tcBorders)
                        
                        # 设置代码块的单元格属性和底纹
                        cell = code_table.cell(0, 0)
                        tcPr = cell._tc.get_or_add_tcPr()
                        
                        # 设置底纹
                        shd_str = '<w:shd {}/>'.format(nsdecls('w'))
                        shd = parse_xml(shd_str)
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), 'F5F5F5')  # 浅灰色底纹
                        tcPr.append(shd)
                        
                        # 使用单元格中的段落添加代码内容
                        code_para = cell.paragraphs[0]
                        code_para.paragraph_format.left_indent = Inches(0.3)
                        code_para.paragraph_format.space_before = Pt(6)
                        code_para.paragraph_format.space_after = Pt(6)
                        
                        # 添加语言标记（如果有）
                        if language:
                            lang_run = code_para.add_run(f"{language}\n")
                            lang_run.bold = True
                            lang_run.font.size = Pt(9)
                            lang_run.font.color.rgb = RGBColor(100, 100, 100)

                        # 添加代码内容
                        code_text = "\n".join(code_block_content)
                        code_run = code_para.add_run(code_text)
                        code_run.font.name = 'Courier New'  # 等宽字体
                        code_run.font.size = Pt(10)  # 小一号的字体大小
                        code_run.font.color.rgb = RGBColor(80, 80, 80)  # 深灰色

                        i += 1  # 跳过结束标记行
                        continue

                # 如果在代码块内，收集内容
                if in_code_block:
                    code_block_content.append(line)
                    i += 1
                    continue

                # 处理分隔线（如 ---、***、___ 独占一行）
                if re.match(r'^[-*_]{3,}$', line):
                    # 添加一个水平线
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Inches(0)
                    para.paragraph_format.right_indent = Inches(0)
                    para.paragraph_format.space_before = Pt(10)
                    para.paragraph_format.space_after = Pt(10)

                    # 添加自定义横线
                    run = para.add_run()
                    # 这里用一个特殊技巧：使用下划线字符加上空格来模拟水平线
                    run.add_text('_' * 70)
                    run.font.color.rgb = RGBColor(200, 200, 200)  # 浅灰色

                    i += 1
                    continue

                # 处理标题 - 正确对应Markdown标题级别和Word标题级别
                heading_match = re.match(r'^(#+)\s+(.*)', line)
                if heading_match:
                    # 获取#的数量作为标题级别
                    hash_count = len(heading_match.group(1))
                    # Markdown中#对应一级标题，##对应二级标题，以此类推
                    level = hash_count
                    # 清理标题文本中可能存在的加粗标记
                    text = re.sub(r'\*\*(.*?)\*\*', r'\1', heading_match.group(2))
                    # 清理双引号
                    text = re.sub(r'"([^"]*)"', r'\1', text)

                    # 添加对应级别的标题
                    heading = doc.add_heading(text, level=level)

                    # 确保标题使用黑体，加粗，不斜体，颜色为黑色
                    for run in heading.runs:
                        run.font.name = '黑体'
                        # 确保中文字体名称可以识别
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        run.bold = True  # 确保加粗
                        run.italic = False
                        run.font.color.rgb = BLACK_COLOR  # 使用常量设置为黑色

                    i += 1
                    continue

                # 处理无序列表
                list_match = re.match(r'^\s*([-*+])\s+(.*)', line)
                if list_match and not self.remove_unordered_list:
                    list_item = doc.add_paragraph(style='List Bullet')

                    # 提取列表项文本并处理其中的加粗
                    text = list_match.group(2)
                    self._add_text_with_bold_and_quotes(list_item, text)

                    i += 1
                    continue
                elif list_match and self.remove_unordered_list:
                    # 如果设置了移除无序列表，则作为普通段落处理
                    para = doc.add_paragraph()
                    self._add_text_with_bold_and_quotes(para, list_match.group(2))
                    i += 1
                    continue

                # 处理有序列表
                ordered_list_match = re.match(r'^\s*(\d+\.)\s+(.*)', line)
                if ordered_list_match and not self.remove_ordered_list:
                    list_item = doc.add_paragraph(style='List Number')
                    self._add_text_with_bold_and_quotes(list_item, ordered_list_match.group(2))
                    i += 1
                    continue
                elif ordered_list_match and self.remove_ordered_list:
                    # 如果设置了移除有序列表，则作为普通段落处理
                    para = doc.add_paragraph()
                    self._add_text_with_bold_and_quotes(para, ordered_list_match.group(2))
                    i += 1
                    continue

                # 处理键值对格式（例如"软件工具: Flask 2.3.2 + SQLAlchemy 2.0"）
                kv_match = re.match(r'^([^:]+):\s*(.*)', line)
                if kv_match and not line.startswith('>'):
                    key = kv_match.group(1).strip()
                    value = kv_match.group(2).strip()

                    # 清理键名中的双引号
                    key = re.sub(r'"([^"]*)"', r'\1', key)

                    # 创建段落
                    para = doc.add_paragraph()

                    # 添加键名（粗体）
                    key_run = para.add_run(f"{key}: ")
                    key_run.bold = True
                    key_run.font.name = '黑体'
                    key_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

                    # 添加值（正常字体），支持格式
                    self._add_text_with_bold_and_quotes(para, value)

                    i += 1
                    continue

                # 处理引用块（以>开头的行）
                if line.startswith('>'):
                    # 使用特殊函数处理引用块及其嵌套内容
                    i = self._process_nested_blockquote(doc, lines, i)
                    continue

                # 处理表格（简化处理，表格转换为制表符）
                if '|' in line and not line.startswith('>'):
                    # 如果启用了表格转Word表格选项并且不是清洁模式
                    if self.table_to_word and not self.table_clean and self.table_conversion == "无":
                        # 提取完整表格
                        table_data, has_header, i = self._extract_table(lines, i)
                        # 处理表格为Word表格
                        self._process_table(doc, table_data, has_header)
                        continue

                    # 检查是否是表格分隔行
                    if re.match(r'^\s*\|?[\s\-:|]+\|?\s*$', line):
                        i += 1
                        continue

                    # 按照表格转换设置处理
                    cells = [cell.strip() for cell in line.strip('|').split('|')]

                    # 根据表格转换选项处理
                    if self.table_clean:
                        # 如果启用了table_clean，直接处理为简单文本，使用空格分隔
                        text = " ".join(cells)
                    elif self.table_conversion == "无":
                        # 默认使用空格
                        text = " ".join(cells)
                    elif self.table_conversion == "空格":
                        text = "    ".join(cells)
                    elif self.table_conversion == "/t":
                        text = "\t".join(cells)
                    elif self.table_conversion == ",":
                        text = ",".join(cells)
                    else:
                        # 默认使用空格
                        text = " ".join(cells)

                    para = doc.add_paragraph()
                    self._add_text_with_bold_and_quotes(para, text)
                    i += 1
                    continue

                # 处理普通段落和其他格式
                para = doc.add_paragraph()

                # 处理文本，保留加粗和引号格式
                self._add_text_with_bold_and_quotes(para, line)
                i += 1

            # 保存Word文档
            doc.save(filepath)

            # 生成处理选项的摘要
            options_summary = []
            if self.remove_italic:
                options_summary.append("去除斜体")
            if self.remove_strikethrough:
                options_summary.append("去除删除线")
            if self.remove_highlight:
                options_summary.append("去除高亮")
            if self.remove_links:
                options_summary.append("去除链接")
            if self.remove_unordered_list:
                options_summary.append("清洗无序列表")
            if self.remove_ordered_list:
                options_summary.append("清洗有序列表")
            if self.table_clean:
                options_summary.append("表格清洁")
            elif self.table_conversion != "无":
                options_summary.append(f"表格转换: {self.table_conversion}")
            if self.table_to_word and not self.table_clean and self.table_conversion == "无":
                options_summary.append("表格转为Word表格")

            options_text = "，".join(options_summary) if options_summary else "保留所有格式"
            success_message = f"文档已成功保存至:\n{filepath}\n\n应用的处理选项: {options_text}"

            self.show_message_popup(success_message)

        except Exception as e:
            error_msg = f"Word导出失败: {str(e)}"
            self.show_message_popup(error_msg)

    def _add_text_with_bold_and_quotes(self, paragraph, text):
        """处理文本，支持加粗、引号、行内代码等格式"""
        # 改进：优先处理嵌套格式
        processed_text = text
        
        # 处理md``内嵌**加粗**的情况
        md_nested_bold_matches = list(re.finditer(r'md``(.*?\*\*.*?\*\*.*?)``', processed_text, re.DOTALL))
        for match in reversed(md_nested_bold_matches):
            # 提取内容并先处理内部的加粗格式
            inner_content = match.group(1)
            # 将内部的加粗格式暂时替换为特殊标记
            inner_content = re.sub(r'\*\*(.*?)\*\*', r'__BOLD__\1__BOLD__', inner_content)
            # 替换回原文本
            start, end = match.span()
            processed_text = processed_text[:start] + inner_content + processed_text[end:]
        
        # 处理**加粗**内嵌md``的情况
        bold_nested_md_matches = list(re.finditer(r'\*\*(.*?md``.*?``.*?)\*\*', processed_text, re.DOTALL))
        for match in reversed(bold_nested_md_matches):
            # 提取内容并先处理内部的md``格式
            inner_content = match.group(1)
            # 将内部的md``格式暂时替换为特殊标记
            inner_content = re.sub(r'md``(.*?)``', r'__MD__\1__MD__', inner_content, flags=re.DOTALL)
            # 替换回原文本
            start, end = match.span()
            processed_text = processed_text[:start] + inner_content + processed_text[end:]
        
        # 处理行内代码 `code`
        inline_code_matches = list(re.finditer(r'`([^`]+)`', processed_text))
        for match in reversed(inline_code_matches):
            # 提取代码内容
            code_content = match.group(1)
            # 将代码替换为特殊标记
            start, end = match.span()
            processed_text = processed_text[:start] + "__CODE__" + code_content + "__CODE__" + processed_text[end:]
        
        # 现在开始正常的格式处理
        # 查找所有加粗部分、双引号包裹的文本、md``包裹的文本和行内代码
        bold_parts = re.finditer(r'\*\*(.*?)\*\*', processed_text)
        quote_parts = re.finditer(r'"([^"]*)"', processed_text)
        md_backticks_parts = re.finditer(r'md``(.*?)``', processed_text, re.DOTALL)  # 添加DOTALL标志支持多行
        code_parts = re.finditer(r'__CODE__(.*?)__CODE__', processed_text)  # 处理行内代码
        
        # 处理特殊标记
        special_bold_parts = re.finditer(r'__BOLD__(.*?)__BOLD__', processed_text)
        special_md_parts = re.finditer(r'__MD__(.*?)__MD__', processed_text)

        # 合并所有匹配并按位置排序
        all_matches = []
        for match in bold_parts:
            all_matches.append((match.span(), 'bold', match.group(1)))
        for match in quote_parts:
            all_matches.append((match.span(), 'quote', match.group(1)))
        for match in md_backticks_parts:
            all_matches.append((match.span(), 'md_backticks', match.group(1)))
        for match in special_bold_parts:
            all_matches.append((match.span(), 'bold', match.group(1)))
        for match in special_md_parts:
            all_matches.append((match.span(), 'md_backticks', match.group(1)))
        for match in code_parts:
            all_matches.append((match.span(), 'code', match.group(1)))

        # 按开始位置排序
        all_matches.sort(key=lambda x: x[0][0])

        last_end = 0
        has_formatting = False

        for (start, end), match_type, content in all_matches:
            has_formatting = True
            # 添加格式前的普通文本
            if start > last_end:
                normal_run = paragraph.add_run(text[last_end:start])
                normal_run.font.name = '宋体'
                normal_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 根据类型添加不同格式的文本
            if match_type == 'bold':
                # 加粗文本使用黑体，加粗
                formatted_run = paragraph.add_run(content)
                formatted_run.bold = True
                formatted_run.font.name = '黑体'
                formatted_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            elif match_type == 'quote':
                # 双引号文本使用蓝色
                formatted_run = paragraph.add_run(f'"{content}"')
                formatted_run.font.name = '宋体'
                formatted_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                formatted_run.font.color.rgb = BLUE_COLOR
            elif match_type == 'md_backticks':
                # md``文本处理为加粗样式
                formatted_run = paragraph.add_run(content)
                formatted_run.bold = True
                formatted_run.font.name = '黑体'
                formatted_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            elif match_type == 'code':
                # 行内代码使用等宽字体和灰色背景
                formatted_run = paragraph.add_run(content)
                formatted_run.font.name = 'Courier New'
                formatted_run.font.size = Pt(10)
                formatted_run.font.color.rgb = RGBColor(80, 80, 80)
                
                # 使用适当的方式添加阴影/底纹效果
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                
                # 确保run有rPr元素
                if formatted_run._element.rPr is None:
                    formatted_run._element.get_or_add_rPr()
                    
                # 添加底纹效果
                shd_str = '<w:shd {}/>'.format(nsdecls('w'))
                shd = parse_xml(shd_str)
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F5F5F5')  # 灰色背景
                formatted_run._element.rPr.append(shd)

            last_end = end

        # 如果没有格式化部分，直接添加整个文本
        if not has_formatting and last_end == 0:
            # 处理其他可能的Markdown标记
            processed_text = self._process_markdown_inline(text)
            normal_run = paragraph.add_run(processed_text)
            normal_run.font.name = '宋体'
            # 添加东亚字体设置
            normal_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            return

        # 添加最后剩余的普通文本
        if last_end < len(text):
            # 移除可能存在的其他Markdown标记
            remaining_text = text[last_end:]
            processed_text = self._process_markdown_inline(remaining_text)

            normal_run = paragraph.add_run(processed_text)
            normal_run.font.name = '宋体'
            # 添加东亚字体设置
            normal_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _process_markdown_inline(self, text):
        """处理内联Markdown标记"""
        processed_text = text
        
        # 处理行内代码
        processed_text = re.sub(r'`([^`]+)`', r'\1', processed_text)
        
        # 改进：处理嵌套格式
        # 处理md``内嵌**加粗**的情况
        md_nested_bold_matches = list(re.finditer(r'md``(.*?\*\*.*?\*\*.*?)``', processed_text, re.DOTALL))
        for match in md_nested_bold_matches:
            inner_content = match.group(1)
            # 清理内部的加粗格式
            inner_content = re.sub(r'\*\*(.*?)\*\*', r'\1', inner_content)
            full_match = match.group(0)
            # 替换整个匹配为处理后的内容
            processed_text = processed_text.replace(full_match, inner_content)
        
        # 处理**加粗**内嵌md``的情况
        bold_nested_md_matches = list(re.finditer(r'\*\*(.*?md``.*?``.*?)\*\*', processed_text, re.DOTALL))
        for match in bold_nested_md_matches:
            inner_content = match.group(1)
            # 清理内部的md``格式
            inner_content = re.sub(r'md``(.*?)``', r'\1', inner_content, flags=re.DOTALL)
            full_match = match.group(0)
            # 替换整个匹配为处理后的内容
            processed_text = processed_text.replace(full_match, inner_content)
        
        # 移除md``格式标记，支持多行
        processed_text = re.sub(r'md``(.*?)``', r'\1', processed_text, flags=re.DOTALL)
        
        # 移除斜体（如果启用了该选项）
        if self.remove_italic:
            processed_text = re.sub(r'(?<!\*)\*(?!\*)(.*?)\*(?!\*)', r'\1', processed_text)
            processed_text = re.sub(r'(?<!_)_(?!_)(.*?)_(?!_)', r'\1', processed_text)
        
        # 移除删除线（如果启用了该选项）
        if self.remove_strikethrough:
            processed_text = re.sub(r'~~(.*?)~~', r'\1', processed_text)
        
        # 移除高亮（如果启用了该选项）
        if self.remove_highlight:
            processed_text = re.sub(r'==(.+?)==', r'\1', processed_text)
        
        # 移除链接（如果启用了该选项）
        if self.remove_links:
            processed_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', processed_text)
        
        return processed_text

    def _extract_table(self, lines, start_index):
        """从Markdown文本中提取完整的表格并返回表格数据和新的索引"""
        table_data = []
        i = start_index
        has_header_separator = False
        
        while i < len(lines):
            line = lines[i].strip()
            
            # 如果行为空或不包含|，则表格结束
            if not line or '|' not in line:
                break
                
            # 检查是否为分隔行（包含 ----|----）
            if re.match(r'^\s*\|?[\s\-:|]+\|?\s*$', line):
                has_header_separator = True
                i += 1
                continue
                
            # 提取单元格数据
            cells = [cell.strip() for cell in line.strip('|').split('|')]
            table_data.append(cells)
            
            i += 1
        
        return table_data, has_header_separator, i
        
    def _process_table(self, doc, table_data, has_header):
        """将提取的表格数据转换为Word表格，支持单元格内的复杂格式"""
        if not table_data or not table_data[0]:
            return
            
        # 创建表格
        rows_count = len(table_data)
        cols_count = max(len(row) for row in table_data)
        
        table = doc.add_table(rows=rows_count, cols=cols_count)
        table.style = 'Table Grid'  # 应用表格网格样式
        
        # 填充表格内容
        for i, row_data in enumerate(table_data):
            for j, cell_content in enumerate(row_data):
                if j < cols_count:  # 确保不超出列数
                    cell = table.cell(i, j)
                    # 处理单元格中的Markdown格式
                    para = cell.paragraphs[0]
                    
                    # 检查单元格内容是否为嵌套格式
                    if cell_content.strip().startswith('>'):
                        # 单元格包含引用块
                        content = cell_content.strip()[1:].strip()
                        para.paragraph_format.left_indent = Inches(0.1)
                        # 使用浅灰背景标识引用块
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        
                        # 创建shd元素
                        shd_str = '<w:shd {}/>'.format(nsdecls('w'))
                        shd = parse_xml(shd_str)
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), 'F0F8FF')  # 浅蓝色背景
                        
                        tcPr.append(shd)
                        # 添加内容
                        self._add_text_with_bold_and_quotes(para, content)
                    elif cell_content.strip().startswith('```') and cell_content.strip().endswith('```'):
                        # 单元格包含代码块
                        # 提取代码块内容和语言（如果有）
                        lines = cell_content.strip().split('\n')
                        language = ""
                        if len(lines) > 0 and lines[0].startswith('```'):
                            language = lines[0][3:].strip()
                        # 提取代码内容（去除开始和结束标记）
                        code_content = '\n'.join(lines[1:-1] if len(lines) > 2 else [])
                        
                        # 如果有语言信息，添加语言标记
                        if language:
                            lang_run = para.add_run(f"{language}\n")
                            lang_run.bold = True
                            lang_run.font.size = Pt(9)
                            lang_run.font.color.rgb = RGBColor(100, 100, 100)
                        
                        # 添加代码内容
                        code_run = para.add_run(code_content)
                        code_run.font.name = 'Courier New'
                        code_run.font.size = Pt(9)
                        code_run.font.color.rgb = RGBColor(80, 80, 80)
                        
                        # 为整个单元格设置背景色
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        
                        # 创建shd元素
                        shd_str = '<w:shd {}/>'.format(nsdecls('w'))
                        shd = parse_xml(shd_str)
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), 'F5F5F5')  # 浅灰色背景
                        
                        tcPr.append(shd)
                    elif '`' in cell_content and not cell_content.strip().startswith('```'):
                        # 处理包含行内代码的单元格
                        self._add_text_with_bold_and_quotes(para, cell_content)
                    else:
                        # 使用现有的格式处理函数
                        self._add_text_with_bold_and_quotes(para, cell_content)
                
                # 如果有表头且这是第一行，设置加粗和居中
                if has_header and i == 0:
                    for run in para.runs:
                        run.bold = True
                        run.font.name = '黑体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # 为表头单元格添加背景色
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    
                    # 创建shd元素
                    shd_str = '<w:shd {}/>'.format(nsdecls('w'))
                    shd = parse_xml(shd_str)
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'EEEEEE')  # 浅灰色背景
                    
                    tcPr.append(shd)
                    
        # 添加表格后的空行
        doc.add_paragraph()

    def _process_nested_blockquote(self, doc, lines, start_index):
        """处理嵌套的引用块，支持引用块中的代码块和列表等复杂格式"""
        quote_content = []
        i = start_index
        
        # 收集引用块内容
        while i < len(lines):
            line = lines[i].strip()
            
            # 如果不是引用块的行，则结束引用块
            if not line.startswith('>'):
                break
            
            # 移除引用标记并保留后面的内容
            content = line[1:].strip()
            quote_content.append(content)
            i += 1
        
        # 创建统一的引用块容器 - 用表格实现连续背景
        blockquote_table = doc.add_table(rows=1, cols=1)
        blockquote_table.autofit = True
        blockquote_table.style = 'Table Grid'  # 基本样式
        blockquote_table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # 设置引用块边框 - 只保留左侧边框，其他边框隐藏
        for cell in blockquote_table._cells:
            tcPr = cell._tc.get_or_add_tcPr()
            
            # 创建仅包含左边框的XML元素（避免多行字符串格式问题）
            tcBorders = parse_xml(r'<w:tcBorders %s><w:top w:val="nil"/><w:left w:val="single" w:sz="8" w:space="0" w:color="CCCCCC"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>' % nsdecls('w'))
            tcPr.append(tcBorders)
        
        # 设置引用块的单元格属性
        cell = blockquote_table.cell(0, 0)
        cell.width = Inches(5.5)  # 设置合适的宽度
        
        # 为单元格设置左侧边距和背景色
        tcPr = cell._tc.get_or_add_tcPr()
        
        # 添加左侧边距
        tcMar = parse_xml(r'<w:tcMar %s><w:left w:w="300" w:type="dxa"/></w:tcMar>' % nsdecls('w'))
        tcPr.append(tcMar)
        
        # 设置单元格背景色
        shd_str = '<w:shd {}/>'.format(nsdecls('w'))
        shd = parse_xml(shd_str)
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F8F8F8')  # 浅灰色背景
        tcPr.append(shd)
        
        # 获取单元格的段落，如果没有则添加一个
        if not cell.paragraphs:
            cell.add_paragraph()
        quote_para = cell.paragraphs[0]
        
        # 处理引用块内容
        j = 0
        while j < len(quote_content):
            if j > 0:  # 第一行内容已经有段落，其余行需要添加换行符
                quote_para.add_run("\n")
                
            line = quote_content[j]
            
            # 处理引用块中的代码块
            if line.strip().startswith("```"):
                code_content, language, new_j = self._extract_nested_code_block(quote_content, j)
                
                # 在引用块内创建嵌套的代码块表格
                # 使用与独立代码块相同的表格处理方式
                code_table = doc.add_table(rows=1, cols=1)
                code_table.autofit = True
                code_table.style = 'Table Grid'
                
                # 移除表格边框
                for cell in code_table._cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcBorders = parse_xml(r'<w:tcBorders %s><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>' % nsdecls('w'))
                    tcPr.append(tcBorders)
                
                # 设置代码块的单元格属性和底纹
                code_cell = code_table.cell(0, 0)
                tcPr = code_cell._tc.get_or_add_tcPr()
                
                # 设置底纹
                shd_str = '<w:shd {}/>'.format(nsdecls('w'))
                shd = parse_xml(shd_str)
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F5F5F5')  # 浅灰色底纹
                tcPr.append(shd)
                
                # 调整单元格的边距，为引用块内的代码块增加更多缩进
                tcMar = parse_xml(r'<w:tcMar %s><w:left w:w="400" w:type="dxa"/></w:tcMar>' % nsdecls('w'))
                tcPr.append(tcMar)
                
                # 使用单元格中的段落添加代码内容
                code_para = code_cell.paragraphs[0]
                code_para.paragraph_format.space_before = Pt(6)
                code_para.paragraph_format.space_after = Pt(6)
                
                # 添加语言标记（如果有）
                if language:
                    lang_run = code_para.add_run(f"{language}\n")
                    lang_run.bold = True
                    lang_run.font.size = Pt(9)
                    lang_run.font.color.rgb = RGBColor(100, 100, 100)
                
                # 添加代码内容
                code_run = code_para.add_run("\n".join(code_content))
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(10)
                code_run.font.color.rgb = RGBColor(80, 80, 80)
                
                # 回到引用块继续处理后面的内容
                j = new_j
                continue
                
            # 处理引用块中的标题
            heading_match = re.match(r'^(#+)\s+(.*)', line)
            if heading_match:
                hash_count = len(heading_match.group(1))
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', heading_match.group(2))
                text = re.sub(r'"([^"]*)"', r'\1', text)
                
                # 在同一个段落中添加标题样式文本
                heading_run = quote_para.add_run(text)
                heading_run.bold = True
                heading_run.font.size = Pt(13 - hash_count)  # 根据标题级别调整大小
                heading_run.font.name = '黑体'
                heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                
                j += 1
                continue
            
            # 处理引用块中的列表
            list_match = re.match(r'^\s*([-*+])\s+(.*)', line)
            if list_match:
                # 在同一个段落中添加列表项，使用特殊的缩进和项目符号
                bullet_run = quote_para.add_run("• ")
                bullet_run.font.name = 'Symbol'
                
                # 添加列表内容
                content_run = quote_para.add_run(list_match.group(2))
                content_run.font.name = '宋体'
                content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
                j += 1
                continue
                
            # 处理引用块中的有序列表
            ordered_list_match = re.match(r'^\s*(\d+\.)\s+(.*)', line)
            if ordered_list_match:
                # 在同一个段落中添加有序列表项
                number_run = quote_para.add_run(f"{ordered_list_match.group(1)} ")
                
                # 添加列表内容
                content_run = quote_para.add_run(ordered_list_match.group(2))
                content_run.font.name = '宋体'
                content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
                j += 1
                continue
            
            # 处理普通文本行
            self._add_text_with_bold_and_quotes(quote_para, line)
            j += 1
        
        return i

    def _extract_nested_code_block(self, lines, start_index):
        """从嵌套内容中提取代码块"""
        code_content = []
        i = start_index + 1  # 跳过开始标记行
        language = ""
        
        # 提取语言信息
        start_line = lines[start_index].strip()
        if start_line.startswith("```"):
            language = start_line[3:].strip()
        
        while i < len(lines):
            line = lines[i]
            # 如果遇到代码块结束标记，退出循环
            if line.strip() == "```":
                i += 1  # 跳过结束标记行
                break
            code_content.append(line)
            i += 1
        
        return code_content, language, i

    def _process_nested_table(self, doc, table_data, has_header):
        """处理嵌套在其他元素中的表格"""
        if not table_data or not table_data[0]:
            return
        
        # 创建表格，但使用更紧凑的样式
        rows_count = len(table_data)
        cols_count = max(len(row) for row in table_data)
        
        table = doc.add_table(rows=rows_count, cols=cols_count)
        table.style = 'Table Grid'  # 应用表格网格样式
        
        # 填充表格内容
        for i, row_data in enumerate(table_data):
            for j, cell_content in enumerate(row_data):
                if j < cols_count:  # 确保不超出列数
                    cell = table.cell(i, j)
                    # 处理单元格中的Markdown格式
                    para = cell.paragraphs[0]
                    
                    # 特殊处理：检查单元格中是否有引用块
                    if cell_content.strip().startswith('>'):
                        content = cell_content.strip()[1:].strip()
                        para.paragraph_format.left_indent = Inches(0.1)
                        # 使用蓝灰色背景标识引用块
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        
                        # 创建shd元素
                        shd_str = '<w:shd {}/>'.format(nsdecls('w'))
                        shd = parse_xml(shd_str)
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), 'F0F8FF')  # 浅蓝色背景
                        
                        tcPr.append(shd)
                        # 添加内容
                        self._add_text_with_bold_and_quotes(para, content)
                    elif cell_content.strip().startswith('```') and cell_content.strip().endswith('```'):
                        # 单元格包含代码块
                        # 提取代码块内容和语言（如果有）
                        lines = cell_content.strip().split('\n')
                        language = ""
                        if len(lines) > 0 and lines[0].startswith('```'):
                            language = lines[0][3:].strip()
                        # 提取代码内容（去除开始和结束标记）
                        code_content = '\n'.join(lines[1:-1] if len(lines) > 2 else [])
                        
                        # 如果有语言信息，添加语言标记
                        if language:
                            lang_run = para.add_run(f"{language}\n")
                            lang_run.bold = True
                            lang_run.font.size = Pt(9)
                            lang_run.font.color.rgb = RGBColor(100, 100, 100)
                        
                        # 添加代码内容
                        code_run = para.add_run(code_content)
                        code_run.font.name = 'Courier New'
                        code_run.font.size = Pt(9)
                        code_run.font.color.rgb = RGBColor(80, 80, 80)
                        
                        # 为整个单元格设置背景色
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        
                        # 创建shd元素
                        shd_str = '<w:shd {}/>'.format(nsdecls('w'))
                        shd = parse_xml(shd_str)
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), 'F5F5F5')  # 浅灰色背景
                        
                        tcPr.append(shd)
                    elif '`' in cell_content and not cell_content.strip().startswith('```'):
                        # 处理包含行内代码的单元格
                        self._add_text_with_bold_and_quotes(para, cell_content)
                    else:
                        # 使用现有的格式处理函数
                        self._add_text_with_bold_and_quotes(para, cell_content)
                
                # 如果有表头且这是第一行，设置加粗和居中
                if has_header and i == 0:
                    for run in para.runs:
                        run.bold = True
                        run.font.name = '黑体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # 为表头单元格添加背景色
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    
                    # 创建shd元素
                    shd_str = '<w:shd {}/>'.format(nsdecls('w'))
                    shd = parse_xml(shd_str)
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'EEEEEE')  # 浅灰色背景
                    
                    tcPr.append(shd)
                    
        # 添加表格后的空行
        doc.add_paragraph()

    def _create_element(self, name, text=None, attributes=None):
        """创建XML元素，可用于文档格式化"""
        # 使用正确的方式导入和使用OxmlElement
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        # 创建XML元素字符串
        nsptag = name.replace('w:', '')
        xml_str = '<w:{} {}/>'.format(nsptag, nsdecls('w'))
        element = parse_xml(xml_str)
        
        # 添加文本内容
        if text:
            element.text = text
        
        # 添加属性
        if attributes:
            for key, value in attributes.items():
                key_name = key.replace('w:', '')
                element.set(qn('w:{}'.format(key_name)), value)
        
        return element

    def show_message_popup(self, message):
        """显示消息弹窗"""
        content = BoxLayout(orientation='vertical', spacing=10, padding=10)
        popup = Popup(title='提示', content=content, size_hint=(0.6, 0.4))
        
        # 添加消息标签
        content.add_widget(Label(text=message))
        
        # 添加确认按钮
        btn = Button(text='确定', size_hint=(1, 0.3))
        btn.bind(on_press=popup.dismiss)
        content.add_widget(btn)
        
        popup.open()

class MarkdownApp(App):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.tray_icon = None
        self.is_running = True
        
        # 确保图标文件存在
        self.ensure_icon_exists()

    def ensure_icon_exists(self):
        """确保应用程序图标文件存在"""
        icon_path = os.path.join('icons', 'mdword.ico')
        if not os.path.exists(icon_path):
            # 如果图标不存在，创建一个
            try:
                if not os.path.exists('icons'):
                    os.makedirs('icons')
                    
                from PIL import Image, ImageDraw
                # 创建128x128的图像
                img = Image.new('RGBA', (128, 128), color=(255, 255, 255, 0))
                draw = ImageDraw.Draw(img)
                
                # 绘制背景和边框
                draw.rectangle([8, 8, 120, 120], fill=(65, 176, 87), outline=(39, 108, 53), width=2)
                
                # 添加文字"MD"
                # 这里我们用简单的矩形来模拟文字
                draw.rectangle([30, 35, 50, 95], fill=(255, 255, 255), outline=(255, 255, 255), width=2)
                draw.rectangle([60, 35, 98, 55], fill=(255, 255, 255), outline=(255, 255, 255), width=2)
                draw.rectangle([60, 65, 98, 95], fill=(255, 255, 255), outline=(255, 255, 255), width=2)
                draw.rectangle([60, 35, 80, 95], fill=(255, 255, 255), outline=(255, 255, 255), width=2)
                
                # 保存为ICO格式
                img.save(icon_path, format='ICO')
            except Exception as e:
                print(f"创建图标文件失败: {e}")

    def build(self):
        Window.size = (800, 500)  # 优化窗口大小
        Window.bind(on_request_close=self.on_request_close)
        self.setup_tray_icon()
        self.register_hotkey()
        self.title = 'mdword'
        
        # 设置应用图标
        if os.path.exists(os.path.join('icons', 'mdword.ico')):
            self.icon = os.path.join('icons', 'mdword.ico')
            
        return MarkdownTool()

    def setup_tray_icon(self):
        def create_image():
            icon_path = os.path.join('icons', 'mdword.ico')
            if os.path.exists(icon_path):
                try:
                    return Image.open(icon_path)
                except:
                    pass
                
            # 如果加载失败，创建一个简单的图标
            from PIL import Image, ImageDraw
            image = Image.new('RGB', (64, 64), 'white')
            dc = ImageDraw.Draw(image)
            dc.rectangle((16, 16, 48, 48), fill='black')
            return image

        import webbrowser
        menu = Menu(
            MenuItem('打开主界面', lambda: self.restore_window()),
            MenuItem('退出', lambda: self.stop_app()),
            MenuItem('关于项目', lambda: webbrowser.open('https://github.com/fhyxz1/mdword'))
        )
        self.tray_icon = Icon(
            'mdword',
            create_image(),
            menu=menu,
            title="mdword\nN+M快速启动"
        )
        threading.Thread(target=self.tray_icon.run, daemon=True).start()

    def register_hotkey(self):
        def toggle_window():
            def _toggle(dt):
                if Window.visible:
                    Window.hide()
                else:
                    Window.show()
                    Window.raise_window()
            Clock.schedule_once(_toggle)
        keyboard.add_hotkey('N+M', Clock.schedule_once(lambda dt: (Window.show(), Window.raise_window())))

    def on_request_close(self, *args):
        self.show_confirmation()
        return True

    def show_confirmation(self):
        content = BoxLayout(orientation='vertical', spacing=10)
        popup = Popup(title='操作确认', content=content, size_hint=(0.6, 0.3))
        btn_layout = BoxLayout(spacing=10, size_hint_y=0.5)
        btn_min = Button(text='最小化到托盘', on_press=lambda x: self.minimize_app(popup))
        btn_exit = Button(text='退出程序', on_press=lambda x: self.exit_app(popup))
        content.add_widget(Label(text='请选择要执行的操作:'))
        btn_layout.add_widget(btn_min)
        btn_layout.add_widget(btn_exit)
        content.add_widget(btn_layout)
        popup.open()

    def minimize_app(self, popup):
        popup.dismiss()
        Clock.schedule_once(lambda dt: Window.hide())

    def exit_app(self, popup):
        popup.dismiss()
        self.stop_app()

    def restore_window(self):
        Clock.schedule_once(lambda dt: (Window.show(), Window.raise_window()))

    def stop_app(self):
        self.is_running = False
        if self.tray_icon:
            self.tray_icon.stop()
        Window.close()
        App.get_running_app().stop()
        os._exit(0)

if __name__ == '__main__':
    MarkdownApp().run()
